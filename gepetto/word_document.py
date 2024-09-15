from docx import Document

class WordDocument:
    def read(self, file_path: str) -> str:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return "\n".join(full_text)

    def write(self, file_path: str, text: str):
        doc = Document()
        doc.add_paragraph(text)
        doc.save(file_path)
